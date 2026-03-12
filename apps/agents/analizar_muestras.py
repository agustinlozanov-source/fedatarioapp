"""
ANALIZADOR DE MUESTRAS — Actas Constitutivas
Fedatario · Correduría Pública No. 3 · Tamaulipas

Lee todos los .docx en samples/actas/ y extrae:
  1. Configuración de página (tamaño, márgenes, espejo)
  2. Fuentes y tamaños usados
  3. Espaciado entre párrafos y altura de línea
  4. Orden y texto de cada sección / cláusula
  5. Estructura de tablas (accionaria, firmas, etc.)
  6. Encabezados y pies de página
  7. Patrones de negritas vs texto normal
  8. Resumen comparativo entre los 8 archivos

Uso:
    python3 analizar_muestras.py
    python3 analizar_muestras.py --detalle   # muestra párrafos completos

Salida:
    samples/reporte_analisis.json   — datos estructurados para actualizar agentes
    samples/reporte_analisis.md     — informe legible para revisión humana
"""

from __future__ import annotations
import argparse
import json
import os
import re
import sys
from collections import Counter, defaultdict
from pathlib import Path
from typing import Any

try:
    from docx import Document
    from docx.oxml.ns import qn
    from lxml import etree
except ImportError:
    print("Instala dependencias: pip install python-docx lxml")
    sys.exit(1)

# ─────────────────────────────────────────────
# CONSTANTES
# ─────────────────────────────────────────────
SAMPLES_DIR = Path(__file__).parent / "samples" / "actas"
OUTPUT_JSON = Path(__file__).parent / "samples" / "reporte_analisis.json"
OUTPUT_MD   = Path(__file__).parent / "samples" / "reporte_analisis.md"

# Palabras clave para detectar secciones del acta
KEYWORDS_SECCION = {
    "encabezado_numero":      r"(?:póliza|número de instrumento|instrumento número)",
    "encabezado_nombre":      r"(?:acta constitutiva|escritura constitutiva)",
    "comparecencia":          r"(?:ante mí|comparecen|COMPARECEN|en mi fe)",
    "objeto_social":          r"(?:objeto social|OBJETO SOCIAL|objeto de la sociedad)",
    "denominacion":           r"(?:denominación social|DENOMINACIÓN SOCIAL)",
    "domicilio_social":       r"(?:domicilio social|DOMICILIO SOCIAL|domicilio de la sociedad)",
    "duracion":               r"(?:duración|DURACIÓN|por tiempo)",
    "capital_social":         r"(?:capital social|CAPITAL SOCIAL)",
    "tabla_accionaria":       r"(?:accionista|Accionista y Registro Federal)",
    "administracion":         r"(?:administración|ADMINISTRACIÓN|administrador único|gerente)",
    "vigilancia":             r"(?:vigilancia|VIGILANCIA|comisario)",
    "asamblea":               r"(?:asamblea|ASAMBLEA)",
    "clausula_reforma":       r"(?:reformas|reforma de estatutos|modificación)",
    "disolucion":             r"(?:disolución|DISOLUCIÓN|liquidación)",
    "protocolizacion":        r"(?:protocolizo|protocolización|PROTOCOLIZO)",
    "firmas":                 r"(?:Firma\.|Huellas Índices|huellas dactilares)",
    "corredor":               r"(?:CORREDOR PÚBLICO|corredor público)",
}

# ─────────────────────────────────────────────
# EXTRACCIÓN DE CONFIGURACIÓN DE PÁGINA
# ─────────────────────────────────────────────

def extraer_config_pagina(doc: Document) -> dict:
    """Extrae dimensiones y márgenes de la primera sección."""
    try:
        sect = doc.sections[0]
        return {
            "ancho_twips":        sect.page_width,
            "alto_twips":         sect.page_height,
            "ancho_cm":           round(sect.page_width / 914400 * 2.54, 2),
            "alto_cm":            round(sect.page_height / 914400 * 2.54, 2),
            "margen_top_twips":   sect.top_margin,
            "margen_bottom_twips":sect.bottom_margin,
            "margen_left_twips":  sect.left_margin,
            "margen_right_twips": sect.right_margin,
            "margen_top_cm":      round(sect.top_margin / 914400 * 2.54, 2) if sect.top_margin else None,
            "margen_bottom_cm":   round(sect.bottom_margin / 914400 * 2.54, 2) if sect.bottom_margin else None,
            "margen_left_cm":     round(sect.left_margin / 914400 * 2.54, 2) if sect.left_margin else None,
            "margen_right_cm":    round(sect.right_margin / 914400 * 2.54, 2) if sect.right_margin else None,
            "margenes_espejo":    _tiene_margenes_espejo(doc),
            "encabezado_diferente_primera": sect.different_first_page_header_footer,
        }
    except Exception as e:
        return {"error": str(e)}

def _tiene_margenes_espejo(doc: Document) -> bool:
    """Detecta si el documento usa márgenes espejo."""
    settings = doc.settings.element
    W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
    return settings.find(f'{W}mirrorMargins') is not None


# ─────────────────────────────────────────────
# EXTRACCIÓN DE FUENTES Y FORMATOS
# ─────────────────────────────────────────────

def extraer_fuentes(doc: Document) -> dict:
    """Recopila todas las fuentes y tamaños utilizados."""
    fuentes: Counter = Counter()
    tamanos: Counter = Counter()
    interlineas: Counter = Counter()

    for para in doc.paragraphs:
        pPr = para._p.find(qn('w:pPr'))
        if pPr is not None:
            sp = pPr.find(qn('w:spacing'))
            if sp is not None:
                line = sp.get(qn('w:line'))
                if line:
                    interlineas[int(line)] += 1

        for run in para.runs:
            if run.font.name:
                fuentes[run.font.name] += 1
            if run.font.size:
                tamanos[run.font.size] += 1

    # Convertir tamaños de EMU a pt (1 pt = 12700 EMU)
    tamanos_pt = {round(sz / 12700): cnt for sz, cnt in tamanos.items()}

    return {
        "fuentes":             dict(fuentes.most_common(10)),
        "fuente_principal":    fuentes.most_common(1)[0][0] if fuentes else None,
        "tamanos_pt":          tamanos_pt,
        "tamano_principal_pt": max(tamanos_pt, key=tamanos_pt.get) if tamanos_pt else None,
        "interlineas_twips":   dict(interlineas.most_common(5)),
    }


# ─────────────────────────────────────────────
# EXTRACCIÓN DE PÁRRAFOS Y SECCIONES
# ─────────────────────────────────────────────

def extraer_secciones(doc: Document, detalle: bool = False) -> list[dict]:
    """
    Analiza cada párrafo y construye un mapa de secciones del acta.
    Devuelve lista con tipo, texto y formato de cada párrafo.
    """
    resultado = []
    for i, para in enumerate(doc.paragraphs):
        texto = para.text.strip()
        if not texto:
            resultado.append({"idx": i, "tipo": "vacio", "texto": ""})
            continue

        # Detectar centrado
        centrado = False
        pPr = para._p.find(qn('w:pPr'))
        if pPr is not None:
            jc = pPr.find(qn('w:jc'))
            if jc is not None:
                centrado = jc.get(qn('w:val')) in ('center', 'Centre')

        # Detectar negritas
        runs_info = []
        for run in para.runs:
            if run.text.strip():
                runs_info.append({
                    "texto": run.text[:80] if not detalle else run.text,
                    "bold":  bool(run.bold),
                    "font":  run.font.name,
                    "sz_pt": round(run.font.size / 12700) if run.font.size else None,
                })
        todo_bold   = all(r["bold"] for r in runs_info) if runs_info else False
        tiene_bold  = any(r["bold"] for r in runs_info) if runs_info else False

        # Detectar sección por keywords
        tipo_detectado = "parrafo"
        for tipo, patron in KEYWORDS_SECCION.items():
            if re.search(patron, texto, re.IGNORECASE):
                tipo_detectado = tipo
                break

        # Detectar numeración de cláusula (ej: "PRIMERA", "SEGUNDA", "I.-", "1.-")
        es_clausula = bool(re.match(
            r'^(?:[IVX]+\.-|[0-9]+\.-|PRIMERA|SEGUNDA|TERCERA|CUARTA|QUINTA|'
            r'SEXTA|SÉPTIMA|OCTAVA|NOVENA|DÉCIMA|[A-Z]+\s*CLÁUSULA)',
            texto, re.IGNORECASE
        ))
        if es_clausula:
            tipo_detectado = "clausula"

        entrada: dict[str, Any] = {
            "idx":       i,
            "tipo":      tipo_detectado,
            "centrado":  centrado,
            "todo_bold": todo_bold,
            "tiene_bold": tiene_bold,
            "texto_corto": texto[:120],
        }
        if detalle:
            entrada["runs"] = runs_info
            entrada["texto_completo"] = texto

        resultado.append(entrada)

    return resultado


def extraer_orden_secciones(analisis_parrafos: list[dict]) -> list[str]:
    """Devuelve solo los tipos de secciones en orden (sin duplicados consecutivos)."""
    tipos = [p["tipo"] for p in analisis_parrafos if p["tipo"] != "vacio"]
    resultado = []
    for t in tipos:
        if not resultado or resultado[-1] != t:
            resultado.append(t)
    return resultado


# ─────────────────────────────────────────────
# EXTRACCIÓN DE TABLAS
# ─────────────────────────────────────────────

def extraer_tablas(doc: Document) -> list[dict]:
    """Analiza cada tabla del documento."""
    tablas = []
    for i, tabla in enumerate(doc.tables):
        filas = []
        for fila in tabla.rows:
            celdas = [c.text.strip()[:80] for c in fila.cells]
            filas.append(celdas)

        # Detectar tipo de tabla
        primer_texto = " ".join(filas[0]) if filas else ""
        tipo = "desconocida"
        if re.search(r'accionista|Accionista', primer_texto, re.IGNORECASE):
            tipo = "tabla_accionaria"
        elif re.search(r'firma|Firma|huella|Huella', primer_texto, re.IGNORECASE):
            tipo = "tabla_firma"
        elif re.search(r'nombre|Nombre', primer_texto, re.IGNORECASE):
            tipo = "tabla_nombre"

        # Anchos de columnas en twips
        try:
            anchos = [c.width for c in tabla.rows[0].cells] if tabla.rows else []
        except Exception:
            anchos = []

        tablas.append({
            "idx":              i,
            "tipo":             tipo,
            "num_filas":        len(filas),
            "num_columnas":     len(tabla.columns),
            "anchos_col_twips": anchos,
            "primera_fila":     filas[0] if filas else [],
            "contenido":        filas[:5],  # primeras 5 filas
        })
    return tablas


# ─────────────────────────────────────────────
# EXTRACCIÓN DE ENCABEZADOS Y PIES
# ─────────────────────────────────────────────

def extraer_encabezados_pies(doc: Document) -> dict:
    """Extrae texto de encabezados y pies de página."""
    resultado = {"encabezados": [], "pies": []}
    for sect in doc.sections:
        try:
            if sect.header:
                texto = "\n".join(p.text for p in sect.header.paragraphs).strip()
                if texto:
                    resultado["encabezados"].append(texto)
        except Exception:
            pass
        try:
            if sect.footer:
                texto = "\n".join(p.text for p in sect.footer.paragraphs).strip()
                if texto:
                    resultado["pies"].append(texto)
        except Exception:
            pass
    return resultado


# ─────────────────────────────────────────────
# ANÁLISIS DE UN ARCHIVO
# ─────────────────────────────────────────────

def analizar_archivo(path: Path, detalle: bool = False) -> dict:
    """Analiza un solo .docx y retorna datos estructurados."""
    print(f"  → Analizando: {path.name}")
    try:
        doc = Document(str(path))
    except Exception as e:
        return {"archivo": path.name, "error": str(e)}

    parrafos    = extraer_secciones(doc, detalle=detalle)
    tablas      = extraer_tablas(doc)
    num_parrafos = len([p for p in parrafos if p["tipo"] != "vacio"])
    num_chars    = sum(len(p.get("texto_completo", p.get("texto_corto", ""))) for p in parrafos)

    return {
        "archivo":          path.name,
        "num_parrafos":     num_parrafos,
        "num_chars":        num_chars,
        "num_tablas":       len(tablas),
        "config_pagina":    extraer_config_pagina(doc),
        "fuentes":          extraer_fuentes(doc),
        "orden_secciones":  extraer_orden_secciones(parrafos),
        "tablas":           tablas,
        "encabezados_pies": extraer_encabezados_pies(doc),
        "parrafos":         parrafos,
    }


# ─────────────────────────────────────────────
# COMPARACIÓN ENTRE ARCHIVOS
# ─────────────────────────────────────────────

def comparar_archivos(resultados: list[dict]) -> dict:
    """Genera resumen comparativo de todos los archivos."""
    validos = [r for r in resultados if "error" not in r]
    if not validos:
        return {}

    # Márgenes — moda
    margenes = defaultdict(list)
    for r in validos:
        cfg = r.get("config_pagina", {})
        for k in ["margen_top_twips","margen_bottom_twips","margen_left_twips","margen_right_twips"]:
            if cfg.get(k):
                margenes[k].append(cfg[k])
    moda_margenes = {k: Counter(v).most_common(1)[0][0] for k, v in margenes.items()}

    # Fuentes — moda
    fuentes_todas: Counter = Counter()
    for r in validos:
        f = r.get("fuentes", {}).get("fuente_principal")
        if f:
            fuentes_todas[f] += 1

    # Tamaños — moda
    tamanos_todos: Counter = Counter()
    for r in validos:
        t = r.get("fuentes", {}).get("tamano_principal_pt")
        if t:
            tamanos_todos[t] += 1

    # Orden de secciones — más frecuente
    todas_ordenes = [tuple(r["orden_secciones"]) for r in validos if "orden_secciones" in r]
    orden_comun   = list(Counter(todas_ordenes).most_common(1)[0][0]) if todas_ordenes else []

    # Tablas presentes
    tipos_tablas: Counter = Counter()
    for r in validos:
        for t in r.get("tablas", []):
            tipos_tablas[t["tipo"]] += 1

    # Interlineado
    interlineas_todos: Counter = Counter()
    for r in validos:
        il = r.get("fuentes", {}).get("interlineas_twips", {})
        for k, v in il.items():
            interlineas_todos[k] += v

    return {
        "archivos_analizados":    len(validos),
        "archivos_con_error":     len(resultados) - len(validos),
        "fuente_comun":           fuentes_todas.most_common(1)[0][0] if fuentes_todas else None,
        "tamano_pt_comun":        tamanos_todos.most_common(1)[0][0] if tamanos_todos else None,
        "interlineado_comun_twips": interlineas_todos.most_common(1)[0][0] if interlineas_todos else None,
        "margenes_comunes_twips": moda_margenes,
        "orden_secciones_comun":  orden_comun,
        "tipos_tablas_detectados": dict(tipos_tablas),
        "espejo_margenes":        sum(1 for r in validos if r.get("config_pagina",{}).get("margenes_espejo")),
    }


# ─────────────────────────────────────────────
# GENERACIÓN DE REPORTE MARKDOWN
# ─────────────────────────────────────────────

def generar_reporte_md(resultados: list[dict], comparacion: dict) -> str:
    lines = ["# Reporte de Análisis — Actas Constitutivas", ""]

    lines.append("## Resumen comparativo")
    lines.append(f"- Archivos analizados: **{comparacion.get('archivos_analizados', 0)}**")
    lines.append(f"- Fuente principal: **{comparacion.get('fuente_comun')}**")
    lines.append(f"- Tamaño fuente: **{comparacion.get('tamano_pt_comun')} pt**")
    lines.append(f"- Interlineado: **{comparacion.get('interlineado_comun_twips')} twips**")
    lines.append(f"- Márgenes espejo: **{'Sí' if comparacion.get('espejo_margenes', 0) > 0 else 'No'}** ({comparacion.get('espejo_margenes', 0)}/{comparacion.get('archivos_analizados', 0)} archivos)")

    m = comparacion.get("margenes_comunes_twips", {})
    if m:
        lines.append(f"- Márgenes (twips): top={m.get('margen_top_twips')} | bottom={m.get('margen_bottom_twips')} | left={m.get('margen_left_twips')} | right={m.get('margen_right_twips')}")

    lines.append(f"- Tablas detectadas: {comparacion.get('tipos_tablas_detectados', {})}")
    lines.append("")

    lines.append("## Orden de secciones más frecuente")
    for i, sec in enumerate(comparacion.get("orden_secciones_comun", []), 1):
        lines.append(f"  {i}. `{sec}`")
    lines.append("")

    lines.append("## Detalle por archivo")
    for r in resultados:
        lines.append(f"### {r['archivo']}")
        if "error" in r:
            lines.append(f"⚠️ Error: {r['error']}")
            lines.append("")
            continue

        cfg = r.get("config_pagina", {})
        fnt = r.get("fuentes", {})
        lines.append(f"- Párrafos: {r.get('num_parrafos')} | Caracteres: {r.get('num_chars')} | Tablas: {r.get('num_tablas')}")
        lines.append(f"- Página: {cfg.get('ancho_cm')} × {cfg.get('alto_cm')} cm")
        lines.append(f"- Márgenes (cm): top={cfg.get('margen_top_cm')} | bottom={cfg.get('margen_bottom_cm')} | left={cfg.get('margen_left_cm')} | right={cfg.get('margen_right_cm')}")
        lines.append(f"- Espejo: {'✓' if cfg.get('margenes_espejo') else '✗'}")
        lines.append(f"- Fuente: {fnt.get('fuente_principal')} | Tamaño: {fnt.get('tamano_principal_pt')} pt")
        lines.append(f"- Interlineados (twips): {fnt.get('interlineas_twips', {})}")

        enc = r.get("encabezados_pies", {})
        if enc.get("encabezados"):
            lines.append(f"- Encabezado: `{enc['encabezados'][0][:80]}`")
        if enc.get("pies"):
            lines.append(f"- Pie: `{enc['pies'][0][:80]}`")

        lines.append(f"- Orden de secciones: `{' → '.join(r.get('orden_secciones', [])[:15])}`")

        for t in r.get("tablas", []):
            lines.append(f"  - Tabla [{t['tipo']}]: {t['num_filas']} filas × {t['num_columnas']} cols | anchos={t['anchos_col_twips']}")
            lines.append(f"    Primera fila: {t['primera_fila']}")

        lines.append("")

    lines.append("---")
    lines.append("*Generado por analizar_muestras.py — Fedatario*")
    return "\n".join(lines)


# ─────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description="Analiza muestras de actas constitutivas .docx")
    parser.add_argument("--detalle", action="store_true", help="Incluir texto completo de cada párrafo")
    args = parser.parse_args()

    archivos = sorted(SAMPLES_DIR.glob("*.docx"))
    if not archivos:
        print(f"\n⚠️  No se encontraron archivos .docx en: {SAMPLES_DIR}")
        print("   Sube los archivos y vuelve a ejecutar.")
        sys.exit(0)

    print(f"\n📂 Encontrados {len(archivos)} archivos en {SAMPLES_DIR.name}/")
    print("=" * 60)

    resultados = [analizar_archivo(f, detalle=args.detalle) for f in archivos]
    comparacion = comparar_archivos(resultados)

    # Guardar JSON
    datos_salida = {"comparacion": comparacion, "archivos": resultados}
    OUTPUT_JSON.write_text(json.dumps(datos_salida, indent=2, ensure_ascii=False, default=str))
    print(f"\n✅ JSON guardado en: {OUTPUT_JSON.name}")

    # Guardar Markdown
    md = generar_reporte_md(resultados, comparacion)
    OUTPUT_MD.write_text(md, encoding="utf-8")
    print(f"✅ Markdown guardado en: {OUTPUT_MD.name}")

    # Resumen en consola
    print("\n" + "=" * 60)
    print("RESUMEN COMPARATIVO")
    print("=" * 60)
    for k, v in comparacion.items():
        print(f"  {k}: {v}")

    print("\n📌 Próximo paso:")
    print("   Comparte reporte_analisis.md con Copilot para actualizar")
    print("   agt04_redactor.py, agt05_auditor.py y agt06_docx.py")


if __name__ == "__main__":
    main()
