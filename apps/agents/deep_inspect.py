"""Inspección profunda: estilos, tablas de firma y tamaño real de fuente."""
from docx import Document
from docx.oxml.ns import qn
from lxml import etree
import sys

def dump_xml(element, indent=0):
    tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag
    attrs = {k.split('}')[-1]: v for k, v in element.attrib.items()}
    print("  " * indent + f"<{tag} {attrs}>")
    if element.text and element.text.strip():
        print("  " * (indent+1) + repr(element.text))
    for child in element:
        dump_xml(child, indent + 1)

# ── 1. Tamaño de fuente desde estilos del documento ─────────────────────────
fname = 'samples/actas/Póliza. SA de CV (1).docx'
doc = Document(fname)
print("=== ESTILOS - tamaños de fuente ===")
for style in doc.styles:
    try:
        rPr = style.element.find('.//' + qn('w:rPr'))
        if rPr is not None:
            sz = rPr.find(qn('w:sz'))
            if sz is not None:
                val = sz.get(qn('w:val'))
                print(f"  Estilo '{style.name}': sz={val} ({int(val)//2}pt)")
    except Exception:
        pass

# ── 2. Tamaño desde docDefaults ─────────────────────────────────────────────
print("\n=== docDefaults ===")
styles_el = doc.styles.element
defaults = styles_el.find(qn('w:docDefaults'))
if defaults is not None:
    dump_xml(defaults)

# ── 3. XML completo de la tabla accionaria real ──────────────────────────────
print("\n=== XML TABLA ACCIONARIA (tabla 0) ===")
tabla = doc.tables[0]
dump_xml(tabla._tbl, indent=0)

# ── 4. XML completo de la tabla de firmas real ───────────────────────────────
print("\n=== XML TABLA FIRMA socio 1 (tabla 1) ===")
tabla_f = doc.tables[1]
dump_xml(tabla_f._tbl, indent=0)

# ── 5. Párrafo inmediatamente antes y después de cada tabla ──────────────────
print("\n=== PÁRRAFOS ALREDEDOR DE LAS TABLAS ===")
body = doc.element.body
children = list(body)
for i, child in enumerate(children):
    tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
    if tag == 'tbl':
        prev = children[i-1] if i > 0 else None
        nxt  = children[i+1] if i < len(children)-1 else None
        print(f"\n  [Tabla encontrada en posición {i}]")
        if prev is not None:
            prev_tag = prev.tag.split('}')[-1]
            prev_txt = ''.join(t.text or '' for t in prev.iter() if t.text)
            print(f"    ANTES ({prev_tag}): {prev_txt[:80]!r}")
        if nxt is not None:
            nxt_tag = nxt.tag.split('}')[-1]
            nxt_txt  = ''.join(t.text or '' for t in nxt.iter()  if t.text)
            print(f"    DESPUÉS ({nxt_tag}): {nxt_txt[:80]!r}")
