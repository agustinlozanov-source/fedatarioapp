"""Script auxiliar para inspección profunda del XML de los .docx"""
from docx import Document
from docx.oxml.ns import qn
from collections import Counter

ARCHIVOS = [
    'samples/actas/Póliza. SA de CV (1).docx',
    'samples/actas/Póliza. SdeRLdeCV 001 (3).docx',
]

for fname in ARCHIVOS:
    print("=== " + fname[-38:] + " ===")
    doc = Document(fname)

    # Tablas
    for t_idx, tabla in enumerate(doc.tables):
        tbl  = tabla._tbl
        grid = tbl.find(qn('w:tblGrid'))
        anchos = [c.get(qn('w:w')) for c in grid.findall(qn('w:gridCol'))] if grid is not None else []
        print("  Tabla %d: %df x %dc | gridCols=%s" % (t_idx, len(tabla.rows), len(tabla.columns), anchos))
        for r_idx, row in enumerate(tabla.rows[:3]):
            celdas = []
            for cell in row.cells:
                tcPr = cell._tc.find(qn('w:tcPr'))
                tcW  = tcPr.find(qn('w:tcW')) if tcPr is not None else None
                w    = tcW.get(qn('w:w'))     if tcW  is not None else None
                celdas.append("w=%s:%s" % (w, cell.text[:22]))
            print("    r%d: %s" % (r_idx, celdas))

    # Interlineado
    lines = []
    for para in doc.paragraphs:
        if not para.text.strip():
            continue
        pPr = para._p.find(qn('w:pPr'))
        if pPr is not None:
            sp = pPr.find(qn('w:spacing'))
            if sp is not None:
                line = sp.get(qn('w:line'))
                if line:
                    lines.append(int(line))
    if lines:
        print("  Interlinea predominante: %s" % Counter(lines).most_common(3))
    print()
